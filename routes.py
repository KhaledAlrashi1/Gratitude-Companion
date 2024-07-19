from flask import Blueprint, request, jsonify, render_template, session, current_app, send_file
import pandas as pd
import logging
import random
from openai import OpenAI
import os
from docx import Document
from io import BytesIO
from docx.shared import RGBColor

main = Blueprint('main', __name__)

# Load OpenAI API key from environment variable
client = OpenAI(
    api_key=os.environ.get("OPENAI_API_KEY"),
)

if not client:
    raise ValueError("The OpenAI API key is not set. Please set the OPENAI_API_KEY environment variable.")


# openai.api_key = os.getenv("OPENAI_API_KEY")

# if not openai.api_key:
#     raise ValueError("The OpenAI API key is not set. Please set the OPENAI_API_KEY environment variable.")


# Load greeting messages from CSV
greeting_messages = pd.read_csv('greeting_messages.csv')['greeting_message'].tolist()

@main.route('/')
def index():
    try:
        logging.info("Fetching initial greeting message from CSV")
        initial_message = random.choice(greeting_messages)
        logging.info(f"Initial message: {initial_message}")
    except Exception as e:
        logging.error(f"Error fetching initial greeting message: {e}")
        initial_message = "Hello! What is something you're grateful for today?"

    return render_template('index.html', initial_message=initial_message)

@main.route('/greeting', methods=['GET'])
def greeting():
    try:
        greeting_message = random.choice(greeting_messages)
        return jsonify({'greeting': greeting_message})
    except Exception as e:
        logging.error(f"Error generating greeting: {e}")
        return jsonify({'error': str(e)}), 500

@main.route('/chat', methods=['POST'])
def chat():
    try:
        data = request.json
        user_input = data.get('message')

        # Retrieve conversation history from the session
        conversation_history = session.get('conversation_history', [])
        conversation_history.append({"role": "user", "content": user_input})

        # Prepare the tailored prompt
        system_message = {
            "role": "system",
            "content": """
            You are a friendly and supportive chatbot focused on helping people practice gratitude.
            
            Praise their thoughtful responses and relate to their answers by sharing similar incidents or 
            expressing understanding.
            
            Give a follow-up question to give a sense of friendly and deep conversation. The follow-up
            questions could be on the same subject or something relevant. Keep the practice exciting 
            and engaging. Maintain a friendly, conversational tone.

            If the user asks something unrelated to gratitude, gently redirect them back to 
            the topic and provide a new gratitude question. Decline any requests for tasks 
            outside the scope of gratitude.

            """
        }

        messages = [system_message] + conversation_history

        logging.info(f"Sending messages to OpenAI model: {messages}")

        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=messages,
            max_tokens=150
        )

        # bot_response = response['choices'][0]['message']['content'].strip()
        bot_response = response.choices[0].message.content

        logging.info(f"Received response from OpenAI: {bot_response}")

        # Append the chatbot's response to the conversation history
        conversation_history.append({"role": "assistant", "content": bot_response})

        # Store the updated conversation history in the session
        session['conversation_history'] = conversation_history

        return jsonify({'response': bot_response})
    except Exception as e:
        logging.error(f"Error during /chat request: {e}")
        return jsonify({'error': str(e)}), 500

@main.route('/export/docx')
def export_docx():
    try:
        conversation_history = session.get('conversation_history', [])
        doc = Document()
        doc.add_heading('Conversation Log', 0)

        for message in conversation_history:
            role = message['role']
            content = message['content']
            paragraph = doc.add_paragraph()
            if role == "user":
                run = paragraph.add_run(f'User: {content}')
                run.bold = True
                run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)
            else:
                run = paragraph.add_run(f'Assistant: {content}')
                run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)

        # Save to a BytesIO stream
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        return send_file(
            file_stream,
            as_attachment=True,
            download_name='conversation.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        logging.error(f"Error during DOCX export: {e}")
        return jsonify({'error': str(e)}), 500